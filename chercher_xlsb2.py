from __future__ import annotations

import argparse
import csv
import sys
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List

try:
    from pyxlsb import open_workbook
except ImportError:  # pragma: no cover - dépend de l'environnement utilisateur
    open_workbook = None

try:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.shared import Cm, Pt, RGBColor
except ImportError:  # pragma: no cover - dépend de l'environnement utilisateur
    Document = None
    WD_ALIGN_VERTICAL = None
    WD_ALIGN_PARAGRAPH = None
    OxmlElement = None
    qn = None
    RT = None
    Cm = None
    Pt = None
    RGBColor = None


COULEUR_PRINCIPALE = "1F4E78"
COULEUR_SECONDAIRE = "D9EAF7"
COULEUR_TEXTE_CLAIR = "FFFFFF"
LIMITE_APERCU_WORD = 250


def numero_vers_colonne_excel(numero: int) -> str:
    """Convertit 1 -> A, 27 -> AA, etc."""
    resultat = ""
    while numero > 0:
        numero, reste = divmod(numero - 1, 26)
        resultat = chr(65 + reste) + resultat
    return resultat


def normaliser_texte(valeur: object) -> str:
    if valeur is None:
        return ""
    return str(valeur).strip()


def correspond(
    texte_cellule: str,
    mot_recherche: str,
    exact: bool = False,
    ignore_case: bool = True,
) -> bool:
    if ignore_case:
        texte_cellule = texte_cellule.casefold()
        mot_recherche = mot_recherche.casefold()

    if exact:
        return texte_cellule == mot_recherche

    return mot_recherche in texte_cellule


def iterer_fichiers_xlsb(racine: Path) -> Iterable[Path]:
    for fichier in racine.rglob("*.xlsb"):
        if not fichier.name.startswith("~$"):
            yield fichier


def apercu_texte(texte: str, limite: int = LIMITE_APERCU_WORD) -> str:
    texte = texte.replace("\r", " ").replace("\n", " ").strip()
    if len(texte) <= limite:
        return texte
    return texte[: limite - 1].rstrip() + "…"


def chemin_vers_url(chemin: Path) -> str:
    """Construit une URL file:// compatible avec les hyperliens Word."""
    absolu = chemin.expanduser().resolve(strict=False)

    try:
        return absolu.as_uri()
    except ValueError:
        texte = str(absolu)
        if texte.startswith("\\\\"):
            return "file://" + texte.lstrip("\\").replace("\\", "/")
        return "file:///" + texte.replace("\\", "/")


def chercher_dans_fichier(
    fichier: Path,
    mot_recherche: str,
    exact: bool = False,
    ignore_case: bool = True,
) -> Iterable[Dict[str, str]]:
    try:
        with open_workbook(str(fichier)) as wb:
            for nom_feuille in wb.sheets:
                try:
                    with wb.get_sheet(nom_feuille) as sheet:
                        for index_ligne, ligne in enumerate(sheet.rows(), start=1):
                            for index_colonne, cell in enumerate(ligne, start=1):
                                valeur = getattr(cell, "v", cell)
                                texte = normaliser_texte(valeur)

                                if not texte:
                                    continue

                                if correspond(texte, mot_recherche, exact=exact, ignore_case=ignore_case):
                                    num_ligne = getattr(cell, "r", None)
                                    num_colonne = getattr(cell, "c", None)

                                    if isinstance(num_ligne, int):
                                        num_ligne += 1
                                    else:
                                        num_ligne = index_ligne

                                    if isinstance(num_colonne, int):
                                        num_colonne += 1
                                    else:
                                        num_colonne = index_colonne

                                    adresse = f"{numero_vers_colonne_excel(num_colonne)}{num_ligne}"

                                    yield {
                                        "fichier": str(fichier.resolve()),
                                        "feuille": str(nom_feuille),
                                        "cellule": adresse,
                                        "valeur": texte,
                                        "erreur": "",
                                    }

                except Exception as e:
                    yield {
                        "fichier": str(fichier.resolve()),
                        "feuille": str(nom_feuille),
                        "cellule": "",
                        "valeur": "",
                        "erreur": f"Erreur lecture feuille: {e}",
                    }

    except Exception as e:
        yield {
            "fichier": str(fichier.resolve()),
            "feuille": "",
            "cellule": "",
            "valeur": "",
            "erreur": f"Erreur ouverture fichier: {e}",
        }


def ecrire_csv(resultats: List[Dict[str, str]], chemin_csv: Path) -> None:
    champs = ["fichier", "feuille", "cellule", "valeur", "erreur"]
    with chemin_csv.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=champs, delimiter=";")
        writer.writeheader()
        writer.writerows(resultats)


def definir_marges(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)


def configurer_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    heading_1 = styles["Heading 1"]
    heading_1.font.name = "Calibri"
    heading_1.font.size = Pt(16)
    heading_1.font.bold = True
    heading_1.font.color.rgb = RGBColor(31, 78, 120)

    heading_2 = styles["Heading 2"]
    heading_2.font.name = "Calibri"
    heading_2.font.size = Pt(12.5)
    heading_2.font.bold = True
    heading_2.font.color.rgb = RGBColor(31, 78, 120)

    heading_3 = styles["Heading 3"]
    heading_3.font.name = "Calibri"
    heading_3.font.size = Pt(11)
    heading_3.font.bold = True
    heading_3.font.color.rgb = RGBColor(54, 96, 146)


def colorer_cellule(cell, couleur_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), couleur_hex)
    tc_pr.append(shd)


def formater_cellule_texte(cell, texte: str, gras: bool = False, couleur: RGBColor | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(texte)
    run.bold = gras
    if couleur is not None:
        run.font.color.rgb = couleur
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def ajouter_hyperlien(paragraph, texte: str, url: str) -> None:
    part = paragraph.part
    relation_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    new_run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = texte
    new_run.append(text)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def tableau_infos(document: Document, titre: str, lignes: List[tuple[str, str]]) -> None:
    document.add_heading(titre, level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False

    entete = table.rows[0].cells
    formater_cellule_texte(entete[0], "Champ", gras=True, couleur=RGBColor(255, 255, 255))
    formater_cellule_texte(entete[1], "Valeur", gras=True, couleur=RGBColor(255, 255, 255))
    colorer_cellule(entete[0], COULEUR_PRINCIPALE)
    colorer_cellule(entete[1], COULEUR_PRINCIPALE)

    entete[0].width = Cm(5)
    entete[1].width = Cm(11)

    for index, (champ, valeur) in enumerate(lignes, start=1):
        cells = table.add_row().cells
        formater_cellule_texte(cells[0], champ, gras=True)
        formater_cellule_texte(cells[1], valeur)
        cells[0].width = Cm(5)
        cells[1].width = Cm(11)
        if index % 2 == 1:
            colorer_cellule(cells[0], "F7FBFF")
            colorer_cellule(cells[1], "F7FBFF")

    document.add_paragraph()


def creer_document_word(
    resultats: List[Dict[str, str]],
    chemin_word: Path,
    mot_recherche: str,
    racine: Path,
    exact: bool,
    case_sensitive: bool,
    nb_fichiers: int,
) -> None:
    if Document is None:
        raise RuntimeError(
            "La librairie python-docx n'est pas installée. Installe-la avec : "
            "python -m pip install python-docx"
        )

    document = Document()
    definir_marges(document)
    configurer_styles(document)

    document.core_properties.title = "Rapport de recherche XLSB"
    document.core_properties.subject = "Résultats de recherche dans des fichiers Excel .xlsb"
    document.core_properties.author = "ChatGPT"

    titre = document.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titre = titre.add_run("Rapport de recherche XLSB")
    run_titre.bold = True
    run_titre.font.name = "Calibri"
    run_titre.font.size = Pt(20)
    run_titre.font.color.rgb = RGBColor(31, 78, 120)

    sous_titre = document.add_paragraph()
    sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sous_titre = sous_titre.add_run(
        f"Mot recherché : {mot_recherche} | Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M:%S')}"
    )
    run_sous_titre.italic = True
    run_sous_titre.font.name = "Calibri"
    run_sous_titre.font.size = Pt(10.5)

    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.add_run(
        "Ce document récapitule les occurrences trouvées dans les fichiers Excel .xlsb. "
        "Les chemins de fichiers affichés ci-dessous sont cliquables depuis Word pour ouvrir directement le fichier correspondant. "
        "Les valeurs très longues sont affichées sous forme d'aperçu pour garder le rapport lisible."
    )

    document.add_paragraph()

    erreurs = [resultat for resultat in resultats if resultat["erreur"]]
    trouves = [resultat for resultat in resultats if not resultat["erreur"]]

    tableau_infos(
        document,
        "Résumé",
        [
            ("Dossier analysé", str(racine)),
            ("Fichiers .xlsb analysés", str(nb_fichiers)),
            ("Occurrences trouvées", str(len(trouves))),
            ("Erreurs détectées", str(len(erreurs))),
            ("Export Word", str(chemin_word)),
        ],
    )

    tableau_infos(
        document,
        "Paramètres de recherche",
        [
            ("Texte recherché", mot_recherche),
            ("Type de correspondance", "Exacte" if exact else "Partielle"),
            ("Sensibilité à la casse", "Oui" if case_sensitive else "Non"),
            ("Format analysé", ".xlsb"),
        ],
    )

    document.add_heading("Résultats par fichier", level=1)

    if not trouves:
        p = document.add_paragraph()
        run = p.add_run("Aucune occurrence trouvée.")
        run.bold = True
        run.font.color.rgb = RGBColor(192, 0, 0)
    else:
        groupes = defaultdict(list)
        for resultat in trouves:
            groupes[resultat["fichier"]].append(resultat)

        for fichier_str in sorted(groupes):
            occurrences = sorted(groupes[fichier_str], key=lambda x: (x["feuille"], x["cellule"], x["valeur"]))
            fichier = Path(fichier_str)

            document.add_heading(fichier.name, level=2)

            p_chemin = document.add_paragraph()
            p_chemin.add_run("Chemin : ").bold = True
            ajouter_hyperlien(p_chemin, fichier_str, chemin_vers_url(fichier))

            p_compte = document.add_paragraph()
            p_compte.add_run("Nombre d'occurrences : ").bold = True
            p_compte.add_run(str(len(occurrences)))

            table = document.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            table.autofit = False

            entete = table.rows[0].cells
            en_tetes = ["Feuille", "Cellule", "Valeur trouvée (aperçu)"]
            largeurs = [Cm(4.5), Cm(2.5), Cm(10.0)]
            for cellule, texte, largeur in zip(entete, en_tetes, largeurs):
                formater_cellule_texte(cellule, texte, gras=True, couleur=RGBColor(255, 255, 255))
                colorer_cellule(cellule, COULEUR_PRINCIPALE)
                cellule.width = largeur

            for index, occ in enumerate(occurrences, start=1):
                cells = table.add_row().cells
                formater_cellule_texte(cells[0], occ["feuille"])
                formater_cellule_texte(cells[1], occ["cellule"], gras=True)
                formater_cellule_texte(cells[2], apercu_texte(occ["valeur"]))
                for cellule, largeur in zip(cells, largeurs):
                    cellule.width = largeur
                if index % 2 == 1:
                    for cellule in cells:
                        colorer_cellule(cellule, "F7FBFF")

            document.add_paragraph()

    if erreurs:
        document.add_heading("Erreurs de lecture", level=1)
        note = document.add_paragraph()
        note.add_run(
            "Les éléments ci-dessous n'ont pas pu être lus entièrement. Vérifie si un fichier est verrouillé, corrompu ou inaccessible."
        )

        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.autofit = False

        entete = table.rows[0].cells
        en_tetes = ["Fichier", "Feuille", "Détail de l'erreur"]
        largeurs = [Cm(6.0), Cm(3.0), Cm(8.0)]
        for cellule, texte, largeur in zip(entete, en_tetes, largeurs):
            formater_cellule_texte(cellule, texte, gras=True, couleur=RGBColor(255, 255, 255))
            colorer_cellule(cellule, "A61C00")
            cellule.width = largeur

        for index, err in enumerate(sorted(erreurs, key=lambda x: (x["fichier"], x["feuille"], x["erreur"])), start=1):
            cells = table.add_row().cells
            p = cells[0].paragraphs[0]
            ajouter_hyperlien(p, err["fichier"], chemin_vers_url(Path(err["fichier"])))
            cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            formater_cellule_texte(cells[1], err["feuille"] or "-")
            formater_cellule_texte(cells[2], err["erreur"])
            for cellule, largeur in zip(cells, largeurs):
                cellule.width = largeur
            if index % 2 == 1:
                colorer_cellule(cells[1], "FFF5F2")
                colorer_cellule(cells[2], "FFF5F2")

    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = footer.add_run("Rapport généré automatiquement")
    run_footer.italic = True
    run_footer.font.size = Pt(9)
    run_footer.font.color.rgb = RGBColor(102, 102, 102)

    document.save(str(chemin_word))


def main() -> int:
    parser = argparse.ArgumentParser(
        description=(
            "Recherche un mot dans tous les fichiers .xlsb d'un dossier et de ses sous-dossiers, "
            "puis génère un CSV et un rapport Word."
        )
    )
    parser.add_argument("mot", help="Mot ou texte à rechercher")
    parser.add_argument(
        "racine",
        nargs="?",
        default=".",
        help="Dossier racine à parcourir (par défaut : dossier courant)",
    )
    parser.add_argument(
        "--exact",
        action="store_true",
        help="Fait une correspondance exacte au lieu d'une recherche partielle",
    )
    parser.add_argument(
        "--case-sensitive",
        action="store_true",
        help="Respecte la casse (majuscules/minuscules)",
    )
    parser.add_argument(
        "--output",
        default="resultats_recherche_xlsb.csv",
        help="Fichier CSV de sortie (par défaut : resultats_recherche_xlsb.csv)",
    )
    parser.add_argument(
        "--word-output",
        default="resultats_recherche_xlsb.docx",
        help="Fichier Word de sortie (par défaut : resultats_recherche_xlsb.docx)",
    )

    args = parser.parse_args()

    if open_workbook is None:
        print(
            "La librairie pyxlsb n'est pas installée. Installe-la avec : python -m pip install pyxlsb",
            file=sys.stderr,
        )
        return 1

    racine = Path(args.racine).expanduser().resolve()
    if not racine.exists() or not racine.is_dir():
        print(f"Le dossier n'existe pas ou n'est pas un dossier : {racine}", file=sys.stderr)
        return 1

    fichiers = list(iterer_fichiers_xlsb(racine))
    if not fichiers:
        print(f"Aucun fichier .xlsb trouvé dans : {racine}")
        return 0

    chemin_csv = Path(args.output).expanduser().resolve()
    chemin_word = Path(args.word_output).expanduser().resolve()

    resultats: List[Dict[str, str]] = []
    nb_fichiers = 0
    nb_occurrences = 0
    nb_erreurs = 0

    for fichier in fichiers:
        nb_fichiers += 1
        for resultat in chercher_dans_fichier(
            fichier=fichier,
            mot_recherche=args.mot,
            exact=args.exact,
            ignore_case=not args.case_sensitive,
        ):
            resultats.append(resultat)

            if resultat["erreur"]:
                nb_erreurs += 1
                print(f"[ERREUR] {resultat['fichier']} | {resultat['feuille']} | {resultat['erreur']}")
            else:
                nb_occurrences += 1
                print(
                    f"[OK] {resultat['fichier']} | "
                    f"Feuille: {resultat['feuille']} | "
                    f"Cellule: {resultat['cellule']} | "
                    f"Valeur: {apercu_texte(resultat['valeur'], 120)}"
                )

    ecrire_csv(resultats, chemin_csv)

    try:
        creer_document_word(
            resultats=resultats,
            chemin_word=chemin_word,
            mot_recherche=args.mot,
            racine=racine,
            exact=args.exact,
            case_sensitive=args.case_sensitive,
            nb_fichiers=nb_fichiers,
        )
        word_message = f"Word généré : {chemin_word}"
    except Exception as e:
        word_message = f"Impossible de générer le Word : {e}"
        print(f"[ERREUR WORD] {e}", file=sys.stderr)

    print("\n--- Résumé ---")
    print(f"Fichiers analysés : {nb_fichiers}")
    print(f"Occurrences trouvées : {nb_occurrences}")
    print(f"Erreurs : {nb_erreurs}")
    print(f"CSV généré : {chemin_csv}")
    print(word_message)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
